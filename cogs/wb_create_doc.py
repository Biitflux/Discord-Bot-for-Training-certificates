import discord
from discord.ext import commands
from discord import app_commands
from sources.logger import CustomLogger,ERROR, INFO, WARNING
from docx import Document
from datetime import date
from database.database import get_kw_data
from sources.infos import bot_picture,def_picture

logger= CustomLogger()

doc = Document("E:\\Discord\\BerechtiBotServer\\sources\\wb_vorlage.docx")
year = date.today()



class WBCreateDoc(commands.Cog):
  def __init__(self, client: commands.Bot):
    self.client = client

  @app_commands.command(name="wb_create_doc", description="Erstellt dir dein Wochenbericht als Word-Datei")
  async def WBCreateDoc_command(self, interaction: discord.Interaction, name: str, kw: int):
    try:
          if interaction.user.id == 266266567157874700:
            db_data = get_kw_data(kw)
            print(db_data)

            placeholders = {
                'STUDEN_NAME': f'{name}',
                'YEAR': f'{year.year}/{kw}',
                'NR': '65',
                'MONDAY-STADLER': f''





              # 'MONDAY-JOCHIM'
              # 'TUESDAY-FOERDER'
              # 'TUESDAY-PLAN'
              # 'TUESDAY-SCHMID'
              # 'WEDNESDAY-ROEMER'
              # 'WEDNESDAY-ICDL'
              # 'WEDNESDAY-SCHMID'
              # 'WEDNESDAY-FOERDER'
              # 'THURSDAY-SPORT'
              # 'THURSDAY-BECK'
              # 'THURSDAY-DEUTSCH'
              # 'THURSDAY-WK'
              # 'THURSDAY-GK'
              # 'FRIDAY-ENGLISCH'
              # 'FRIDAY-ROEMER'
              # 'FRIDAY-BWL'
          }

            for table in doc.tables:
              replace_in_table(table, placeholders)

            doc.save(f"wb_{name}_kw_{kw}.docx")  
            await interaction.response.send_message("Dein Dokument wurde erfolgreich erstellt!", ephemeral=True)     
            logger.log(f"The document of {name} for the calendar week {kw} was successfully created",level=INFO)   

    except Exception as e:
      logger.log(f"An error occurred in WBCreateDoc_command: {str(e)}", level=ERROR)

    else:
      await interaction.response.send_message(embed=embed_no_permmision, ephemeral=True)      


async def setup(client:commands.Bot) -> None:
  await client.add_cog(WBCreateDoc(client))


def replace_in_table(table, placeholders):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    for placeholder, replacement in placeholders.items():
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, replacement)



embed_no_permmision = discord.Embed( title="**Information**", description="", color=discord.Color.red())
embed_no_permmision.set_thumbnail(url=bot_picture)
embed_no_permmision.set_author(name="Developed by Bitflux_",icon_url=def_picture,url="https://discordapp.com/users/266266567157874700")
embed_no_permmision.add_field(name=f"**Dazu hast du keine Berechtigung!**", value="")
logger.log(f"Someone try to use /channel_message Command", level=WARNING)